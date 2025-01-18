import streamlit as st
from frontend.all_utils import switch_page
import io
from docx import Document

async def load_previous_chats_page():
    with st.sidebar:
        if st.button(label="Back to Main Page", use_container_width=True, icon=":material/arrow_back:", type="primary"):
            switch_page(page_name="main")

    current_login_session = st.session_state["login_sessions"]
    all_chats = current_login_session["chat_sessions"]

    filtered_chats = [chat for chat in all_chats if len(chat["chat_history"]) > 1]
    st.header(body="Previous Chats", anchor=False)

    if not filtered_chats:
        st.info(body="No previous chats available.")
        return

    for index, chat in enumerate(iterable=filtered_chats):
        selected_student = chat["selected_student"]["name"]
        chat_summary = f"{index + 1}. Chat with {selected_student}"

        session_key = f"show_chat_{index}"
        if session_key not in current_login_session:
            current_login_session[session_key] = False

        def toggle_chat_history():
            current_login_session[session_key] = not current_login_session[session_key]

        def generate_word_document(chat):
            doc = Document()
            doc.add_heading(text=f"Chat with {chat['selected_student']['name']}", level=1)

            for msg in chat["chat_history"]:
                role = "You" if msg["role"] == "user" else chat['selected_student']['name']
                content = msg["content"]
                doc.add_paragraph(text=f"{role}: {content}")

            buffer = io.BytesIO()
            doc.save(path_or_stream=buffer)
            buffer.seek(0)
            return buffer

        with st.expander(label=chat_summary, expanded=False):
            st.write(f"Started At: {chat['chat_start_timestamp']}")
            st.write(f"Ended At: {chat['chat_end_timestamp']}")
            st.write(f"Total Messages: {len(chat['chat_history'])}")
            st.markdown(body="---")

            st.button(label="Show Chat History" if not current_login_session[session_key] else "Hide Chat History", key=f"chat_button_{index}", on_click=toggle_chat_history)

            word_buffer = generate_word_document(chat=chat)
            st.download_button(label="Download Chat History", data=word_buffer, file_name=f"chat_{index + 1}_history.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            if current_login_session[session_key]:
                st.subheader(body=f"Chat History with {selected_student}", anchor=False)
                for msg in chat["chat_history"]:
                    role = "You" if msg["role"] == "user" else selected_student
                    content = msg["content"]
                    st.markdown(body=f"**{role}:** {content}")
                st.markdown(body="---")